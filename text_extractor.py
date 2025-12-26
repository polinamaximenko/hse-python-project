import os
import re
import tempfile
import logging
import pypdf
import pytesseract
from PIL import Image
from docx import Document
from pdf2image import convert_from_path

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%H:%M:%S',
    )

logger = logging.getLogger(__name__)


class TextExtractor:
    """Класс для извлечения текста из загруженных файлов / инпута"""
    
    def __init__(self, tesseract_cmd_path: str = None, poppler_path: str = None):
        self.poppler_path = poppler_path
        self.tesseract_cmd_path = tesseract_cmd_path
        self.supported_extensions = ['.txt', '.pdf', '.docx', '.doc',
                                    '.jpg', '.jpeg', '.png', '.bmp', '.tiff']
    
    def extract_from_input(self, text: str) -> str:
        """Обработка текста из инпута"""
        logger.info(f"Обработка введенного текста, длина: {len(text)} символов")
        try:
            cleaned_text = self._clean_text(text)
            return cleaned_text
        except Exception as e:
            logger.error(f"Ошибка при обработке введенного текста: {str(e)}")
            raise

    def extract_from_uploaded_file(self, file_storage) -> str:
        """
        Извлечение текста из загруженного файла в Flask
        
        Args:
            file_storage: объект FileStorage из request.files
        """
        try:
            # Получаем имя файла и расширение
            filename = file_storage.filename
            ext = os.path.splitext(filename)[-1].lower()
            if ext not in self.supported_extensions:
                raise ValueError(f"Расширение {ext} не поддерживается")
            
            logger.info(f"Начало обработки файла: {filename}")
            
            # Создаем временный файл и записываем туда содержимое загруженного файла
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
                temp_path = temp_file.name
                file_storage.save(temp_path)
            
            try:
                # Извлекаем текст в зависимости от типа файла
                if ext == '.pdf':
                    text = self._extract_from_pdf(temp_path)
                elif ext in ['.docx', '.doc']:
                    text = self._extract_from_docx(temp_path)
                elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                    text = self._extract_from_image(temp_path)
                elif ext == '.txt':
                    text = self._extract_from_txt(temp_path)
                
                # Очищаем текст
                cleaned_text = self._clean_text(text)
                
                logger.info(f"Успешно обработан файл {filename}, извлечено {len(cleaned_text)} символов")
                return cleaned_text
                
            finally:
                # Удаляем временный файл
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                    
        except Exception as e:
            logger.error(f"Ошибка при обработке файла {filename}: {e}")
            raise


    def extract_from_path(self, file_path: str) -> str:
        """Извлечение текста из файла по пути (для отладки)"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Файл не найден: {file_path}")
            
            filename = os.path.basename(file_path)
            ext = os.path.splitext(filename)[-1].lower()

            if ext not in self.supported_extensions:
                raise ValueError(f"Расширение {ext} не поддерживается")
            
            # Обрабатываем в зависимости от типа
            if ext == '.pdf':
                text = self._extract_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                text = self._extract_from_docx(file_path)
            elif ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']:
                text = self._extract_from_image(file_path)
            
            cleaned_text = self._clean_text(text)
            
            return cleaned_text
            
        except Exception as e:
            logger.error(f"Ошибка при обработке файла {file_path}: {str(e)}")
            raise

    def _extract_from_pdf(self, pdf_path: str) -> str:
        """Извлечение текста из файла PDF с использованием OCR при необходимости"""
        try:
            standard_text = self._extract_from_pdf_standard(pdf_path)

            if standard_text:
                logger.info(f"PDF содержит распознанный текст: {len(standard_text)} символов")
                return standard_text

            else:
                logger.info("Текст в PDF не распознан, применяется OCR...")
                ocr_text = self._extract_from_pdf_with_ocr(pdf_path)

                combined_text = standard_text + "\n" + ocr_text if standard_text else ocr_text
                logger.info(f"После OCR извлечено {len(combined_text)} символов")
                return combined_text

        except Exception as e:
            logger.error(f"Ошибка при обработке PDF: {str(e)}")
            raise

    def _extract_from_pdf_standard(self, pdf_path: str) -> str:
        """Извлечение текста из распознанного PDF файла"""
        text = ""
        with open(pdf_path, 'rb') as file:
            reader = pypdf.PdfReader(file)

            # Проверяем, не зашифрован ли PDF
            if reader.is_encrypted:
                logger.error("PDF файл зашифрован")
                return ""
            
            # Извлекаем текст со всех страниц
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    else:
                        logger.warning(f"На странице {page_num} не найден текст")
                except Exception as e:
                    logger.error(f"Ошибка при обработке страницы {page_num}: {str(e)}")
        
        logger.info(f"Из файла PDF извлечено {len(text)} символов")
        return text

    def _extract_from_pdf_with_ocr(self, pdf_path: str) -> str:
        """Извлечение текста из PDF с помощью OCR"""

        if not self.poppler_path:
            logger.error("Для распознавания текста в PDF укажите путь к poppler/bin")
        if not self.tesseract_cmd_path:
            logger.error("Для распознавания текста в PDF укажите путь к tesseract.exe")
        else:
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd_path

        try:
            # Конвертируем PDF в изображения
            images = convert_from_path(pdf_path, dpi=300, poppler_path=self.poppler_path)
            logger.info(f"PDF конвертирован в {len(images)} изображений для OCR")

            text_parts = []

            # Обрабатываем каждое изображение
            for i, image in enumerate(images):
                try:
                    # Конвертируем в оттенки серого для лучшего распознавания
                    if image.mode != 'L':
                        image = image.convert('L')

                    # Применяем OCR
                    page_text = pytesseract.image_to_string(image, lang='rus+eng')

                    if page_text and page_text.strip():
                        text_parts.append(f"{page_text}\n")
                        logger.debug(f"Страница {i + 1}: распознано {len(page_text)} символов")
                    else:
                        logger.debug(f"Страница {i + 1}: текст не распознан")

                except Exception as e:
                    logger.error(f"Ошибка при OCR страницы {i + 1}: {str(e)}")

            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Ошибка при конвертации PDF в изображения: {str(e)}")
            raise

    def _extract_from_docx(self, docx_path: str) -> str:
        """Извлечение текста из DOCX файла"""
        doc = Document(docx_path)
        text = ""
        
        # Извлекаем текст из всех абзацев
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        
        # Извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + "\n"
        
        logger.info(f"Из файла DOCX извлечено {len(text)} символов")
        return text
    
    def _extract_from_image(self, image_path: str) -> str:
        """Извлечение текста с изображения с помощью OCR"""
        if not self.tesseract_cmd_path:
            logger.error("Для распознавания текста с изображения укажите путь к tesseract.exe")
        else:
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd_path

        # Открываем изображение
        image = Image.open(image_path)
        
        # Используем pytesseract для распознавания текста
        text = pytesseract.image_to_string(image, lang='rus+eng')
        
        logger.info(f"С изображения извлечено {len(text)} символов")
        return text
    
    def _extract_from_txt(self, txt_path: str) -> str:
        """Извлечение текста из текстового файла"""
        # Пробуем разные кодировки
        encodings = ['utf-8', 'cp1251', 'koi8-r', 'iso-8859-1']
        
        for enc in encodings:
            try:
                with open(txt_path, 'r', encoding=enc) as file:
                    text = file.read()
                    logger.info(f"Из файла TXT прочитан текст с кодировкой {enc}")
                    return text
            except UnicodeDecodeError:
                continue
        raise UnicodeDecodeError("Не удалось определить кодировку файла")
    
    def _clean_text(self, text: str) -> str:
        """
        Очистка текста для дальнейшей обработки
        """
        if not text:
            return ""
        
        # Удаляем лишние пробелы
        cleaned_text = re.sub(r'\s+', ' ', text)
        # Сохраняем базовую пунктуацию, удаляем специальные символы
        cleaned_text = re.sub(r'[^\w\s\n.,;:!?%-]', '', cleaned_text)

        logger.info(f"Извлеченный текст: '{cleaned_text[:100]}...{cleaned_text[-100:]}'")
        return cleaned_text.strip()


if __name__ == '__main__':
    extractor = TextExtractor(tesseract_cmd_path="Tesseract-OCR/tesseract.exe",
                              poppler_path="poppler/Library/bin")
    result = extractor.extract_from_path('examples/example1.jpg')